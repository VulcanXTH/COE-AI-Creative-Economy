# -*- coding: utf-8 -*-
"""ประกอบผลลัพธ์จาก workflow → master .md + .docx (BDI format, ฉบับยาว)"""
import io, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

R = lambda f: io.open(f, encoding='utf-8').read()
body = R('_wf_body.md')
execsum = R('_wf_execSummary.md')
refs = R('_wf_references.md')
figs = R('_wf_figures.md')

# ---------- clean ----------
CHATTER = re.compile(r'(ตามที่มอบหมาย|^I[\'’]?ll\b|^I have\b|^I[\'’]?m\b|^I am\b|^Here is\b|^Here[’\']s\b|^Both\b|^The latest\b|^The task\b|^Executive summary\b|^Now I\b|^The orchestrator\b|following the (established )?format|^I[\'’]?ve\b|requested\. I)', re.I)

def strip_chatter(text):
    out=[]
    for ln in text.split('\n'):
        if CHATTER.search(ln.strip()):
            continue
        out.append(ln)
    return '\n'.join(out)

# split body into section blocks, drop the ===== labels, clean chatter
blocks = re.split(r'\n=====\s*S\d+[^\n]*=====\n', '\n'+body)
body_clean = '\n'.join(strip_chatter(b) for b in blocks)

# normalize program/section headers
def norm_headers(t):
    t = re.sub(r'^#{0,4}\s*\(3\)\s*โปรแกรม', '## โปรแกรม', t, flags=re.M)
    # bare "(n) Title" lines (n=3..10) that are short -> ## header
    def repl(m):
        line=m.group(0).strip()
        return '## '+line if len(line)<75 else line
    t = re.sub(r'^\((?:3|4|5|6|7|8|9|10)\)\s+[^\n]+$', repl, t, flags=re.M)
    # standalone governance header
    t = t.replace('\nกลไกการดำเนินงานและธรรมาภิบาล\n', '\n## กลไกการดำเนินงานและธรรมาภิบาล\n')
    return t
body_clean = norm_headers(body_clean)

exec_clean = strip_chatter(execsum)
refs_clean = strip_chatter(refs)

# ---------- assemble master markdown ----------
# split cover (S1 head) from (1)
m = re.search(r'\n##\s*\(1\)\s', body_clean)
cover = body_clean[:m.start()].strip() if m else ''
rest  = body_clean[m.start():].strip() if m else body_clean

master = (cover + '\n\n---\n\n# บทสรุปผู้บริหาร (Executive Summary)\n\n'
          + exec_clean.split('บทสรุปผู้บริหาร (Executive Summary)')[-1].strip()
          + '\n\n---\n\n' + rest
          + '\n\n---\n\n# ภาคผนวก ก — ข้อมูลอ้างอิงและสถิติสนับสนุน\n\n'
          + refs_clean.split('ข้อมูลอ้างอิงและสถิติสนับสนุน',1)[-1].strip()
          + '\n\n---\n\n# ภาคผนวก ข — รายการภาพประกอบ/แผนภาพ (Figures)\n\n'
          + figs.split('รายการภาพประกอบ/แผนภาพ',1)[-1].split('## ร่างโครง SVG')[0].strip())

io.open('59_coe-creative-fullproposal.md','w',encoding='utf-8').write(master)
print('master md len', len(master))

# ---------- build docx ----------
FONT='TH Sarabun New'; BLUE=(0x1F,0x4E,0x78); GRAY=(0x55,0x55,0x55)
doc=Document()
st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(15)
st._element.rPr.rFonts.set(qn('w:cs'),FONT); st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
def cs(run):
    run.font.name=FONT; rf=run._element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn('w:cs'),FONT); rf.set(qn('w:ascii'),FONT); rf.set(qn('w:hAnsi'),FONT)
def add_runs(p, text, base_bold=False, size=15, color=None):
    for i,part in enumerate(text.split('**')):
        if part=='':
            continue
        r=p.add_run(part); r.bold = base_bold or (i%2==1); r.font.size=Pt(size)
        if color: r.font.color.rgb=RGBColor(*color)
        cs(r)
def para(text='', bold=False, size=15, align=None, color=None, after=6):
    p=doc.add_paragraph()
    if align: p.alignment=align
    p.paragraph_format.space_after=Pt(after)
    if text: add_runs(p, text, bold, size, color)
    return p

def emit_table(rows):
    # rows: list of list-of-str ; first row header, second row is separator (skip)
    data=[r for i,r in enumerate(rows) if not re.match(r'^[\s:|-]+$', '|'.join(r))]
    if not data: return
    ncol=max(len(r) for r in data)
    t=doc.add_table(rows=0, cols=ncol); t.style='Table Grid'
    for ri,row in enumerate(data):
        cells=t.add_row().cells
        for ci in range(ncol):
            val=row[ci] if ci<len(row) else ''
            cells[ci].text=''
            pp=cells[ci].paragraphs[0]
            add_runs(pp, val.strip(), base_bold=(ri==0), size=13)
    para('', after=4)

lines=master.split('\n')
i=0; in_cover=True
while i<len(lines):
    ln=lines[i].rstrip()
    s=ln.strip()
    if s.startswith('## (1)'): in_cover=False
    # table detection
    if s.startswith('|') and '|' in s[1:]:
        tbl=[]
        while i<len(lines) and lines[i].strip().startswith('|'):
            cells=[c.strip() for c in lines[i].strip().strip('|').split('|')]
            tbl.append(cells); i+=1
        emit_table(tbl); continue
    if s=='' or s=='---':
        i+=1; continue
    if s.startswith('<div') or s.startswith('</div'):
        i+=1; continue
    al = WD_ALIGN_PARAGRAPH.CENTER if in_cover else None
    if s.startswith('#### ') or s.startswith('### '):
        para(s.lstrip('# ').strip(), bold=True, size=15, color=BLUE, after=3, align=al)
    elif s.startswith('## '):
        para(s[3:].strip(), bold=True, size=16, color=BLUE, after=4, align=al)
    elif s.startswith('# '):
        para(s[2:].strip(), bold=True, size=(20 if in_cover else 17), color=BLUE, after=6, align=al or WD_ALIGN_PARAGRAPH.CENTER)
    elif s.startswith('> '):
        para(s[2:].strip(), size=13, color=GRAY, after=4, align=al)
    elif re.match(r'^[-*•]\s', s):
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Cm(1.0); p.paragraph_format.space_after=Pt(2)
        add_runs(p, re.sub(r'^[-*•]\s','',s), size=14)
    else:
        para(s, after=5, align=al)
    i+=1

doc.save('59_COE_Creative_Economy_fullproposal.docx')
print('saved docx')
