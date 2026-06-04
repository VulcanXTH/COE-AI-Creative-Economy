"""เติม Section 9 (Creative Economy) ลง .docx ของ สดช.
ตาม template มาตรฐาน: 1 ประธาน + 5 ผู้ทรงคุณวุฒิ + 8 หน่วยงาน + 3 เลขา = 17

Mapping จาก v2 (17+4) → template (17):
  • ประธาน         = นายสุธัช เจริญผล (AICAT)
  • ผู้ทรงคุณวุฒิ 5 = สมาคมผู้กำกับ, TACGA, DAAT, Content Creator, อ.สมชาย
  • หน่วยงาน 8     = depa, วธ., CEA, IP, NIA, BDI, DITP, สมาคมโอเพนซอร์ส
  • เลขา 3        = AICAT, CEA, สดช.
  • CGI → ที่ปรึกษา (ไม่ใส่ใน 17)
"""
from docx import Document
from copy import deepcopy
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

INPUT = '/Users/maxvulcanx/Downloads/Tham/690518 (ร่าง) คำสั่งแต่งตั้งคณะทำงาน COE Updated.docx'
OUTPUT = '/Users/maxvulcanx/Desktop/AI Project/COE AI Creative Economy/14_คำสั่งแต่งตั้ง_Creative_Filled.docx'

def set_para_text(para, new_text):
    """Replace paragraph text, preserving first run formatting."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)

def insert_para_after(ref_para, new_text):
    """Insert paragraph after ref_para with same XML style. Returns new Paragraph."""
    new_p = deepcopy(ref_para._p)
    # Clear all text in copy
    for t in new_p.iter(qn('w:t')):
        t.text = ''
    # Set first text element
    first_t = new_p.find('.//' + qn('w:t'))
    if first_t is not None:
        first_t.text = new_text
    ref_para._p.addnext(new_p)
    return Paragraph(new_p, ref_para._parent)

def find_section_bounds(doc, section_keyword):
    """Find paragraph indices of section bounded by keyword."""
    start = None
    for i, para in enumerate(doc.paragraphs):
        if section_keyword in para.text and 'คณะทำงานขับเคลื่อน' in para.text:
            start = i
            break
    if start is None:
        return None, None
    end = len(doc.paragraphs)
    for i in range(start + 1, len(doc.paragraphs)):
        txt = doc.paragraphs[i].text.strip()
        if txt.startswith('คณะทำงานขับเคลื่อน'):
            end = i
            break
    return start, end

# ════════════════ STEP 1: Modify member slots ════════════════
doc = Document(INPUT)
sec_start, sec_end = find_section_bounds(doc, 'อุตสาหกรรมสร้างสรรค์')
print(f'Section 9 (Creative): paragraphs {sec_start} to {sec_end-1}')

# Member slot replacements (old text → new text, within section 9 only)
member_replacements = [
    # ๙.๑ ประธาน
    ('๙.๑ ผู้แทนหน่วยงานที่รับผิดชอบ (รัฐ/เอกชน)',
     '๙.๑ นายสุธัช เจริญผล นายกสมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย (AICAT)'),
    # ๙.๒ — ผู้ทรงคุณวุฒิคนที่ 1 (others added below)
    ('๙.๒ - ๙.๖ ผู้ทรงคุณวุฒิที่ประธานแต่งตั้ง จำนวนไม่เกิน ๕ ท่าน',
     '๙.๒ นายก/ผู้แทน สมาคมผู้กำกับภาพยนตร์ไทย'),
    # ๙.๗-๙.๑๔ หน่วยงาน 8 ราย
    ('๙.๗ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๑)',
     '๙.๗ ผู้แทน สำนักงานส่งเสริมเศรษฐกิจดิจิทัล (depa)'),
    ('๙.๘ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๒)',
     '๙.๘ ผู้แทน กรมส่งเสริมวัฒนธรรม กระทรวงวัฒนธรรม'),
    ('๙.๙ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๓)',
     '๙.๙ ผู้อำนวยการ/ผู้แทน สำนักงานส่งเสริมเศรษฐกิจสร้างสรรค์ (CEA)'),
    ('๙.๑๐ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๔)',
     '๙.๑๐ อธิบดี/ผู้แทน กรมทรัพย์สินทางปัญญา กระทรวงพาณิชย์'),
    ('๙.๑๑ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๕)',
     '๙.๑๑ ผู้อำนวยการ/ผู้แทน สำนักงานนวัตกรรมแห่งชาติ (NIA)'),
    ('๙.๑๒ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๖)',
     '๙.๑๒ ผู้อำนวยการ/ผู้แทน สถาบันข้อมูลขนาดใหญ่ (BDI)'),
    ('๙.๑๓ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๗)',
     '๙.๑๓ อธิบดี/ผู้แทน กรมส่งเสริมการค้าระหว่างประเทศ (DITP)'),
    ('๙.๑๔ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๘)',
     '๙.๑๔ นายก/ผู้แทน สมาคมสมาพันธ์โอเพนซอร์สแห่งประเทศไทย'),
    # ๙.๑๕ เลขานุการ + ๙.๑๖ ผู้ช่วยเลขา
    ('๙.๑๕ ผู้แทนหน่วยงานที่รับผิดชอบ',
     '๙.๑๕ ผู้แทน สมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย (AICAT)'),
    ('๙.๑๖ ผู้แทนหน่วยงานที่รับผิดชอบ',
     '๙.๑๖ ผู้แทน สำนักงานส่งเสริมเศรษฐกิจสร้างสรรค์ (CEA)'),
]

# Apply within section 9 only
for i in range(sec_start, sec_end):
    para = doc.paragraphs[i]
    for old, new in member_replacements:
        if old in para.text:
            set_para_text(para, para.text.replace(old, new))
            break

# ════════════════ STEP 2: Insert ๙.๓-๙.๖ after ๙.๒ ════════════════
ref_para = None
for i in range(sec_start, sec_end):
    if doc.paragraphs[i].text.strip().startswith('๙.๒ '):
        ref_para = doc.paragraphs[i]
        break

if ref_para:
    inserts = [
        '๙.๓ นายก/ผู้แทน สมาคมผู้ประกอบการแอนิเมชันและคอมพิวเตอร์กราฟิกไทย (TACGA)\t\t\tคณะทำงาน',
        '๙.๔ นายก/ผู้แทน สมาคมโฆษณาดิจิทัล (DAAT)\t\t\t\t\tคณะทำงาน',
        '๙.๕ นายก/ผู้แทน สมาคม Content Creator\t\t\t\t\tคณะทำงาน',
        '๙.๖ ผู้ทรงคุณวุฒิด้านกฎหมายทรัพย์สินทางปัญญาและปัญญาประดิษฐ์\t\tคณะทำงาน',
    ]
    last = ref_para
    for t in inserts:
        last = insert_para_after(last, t)
    print(f'  + Inserted {len(inserts)} ผู้ทรงคุณวุฒิ paragraphs')

# Save intermediate to refresh paragraph indices
doc.save(OUTPUT)

# ════════════════ STEP 3: Fill duties ๒-๕ (reload to refresh) ════════════════
doc = Document(OUTPUT)
sec_start, sec_end = find_section_bounds(doc, 'อุตสาหกรรมสร้างสรรค์')

duty_replacements = {
    '๒.': '๒. ขับเคลื่อนการประยุกต์ใช้ปัญญาประดิษฐ์ในอุตสาหกรรมสร้างสรรค์ของประเทศ ครอบคลุมสาขาภาพยนตร์ แอนิเมชัน ดนตรี โฆษณา คอนเทนต์ดิจิทัล และผู้สร้างสรรค์เนื้อหา รวมถึงพัฒนาและสนับสนุนโครงการต้นแบบ (Use Cases) ให้เกิดผลเป็นรูปธรรมและขยายผลเชิงพาณิชย์',
    '๓.': '๓. สนับสนุนการสร้างชุดข้อมูล Thai Creative AI Dataset โครงสร้างพื้นฐานข้อมูล และเครื่องมือปัญญาประดิษฐ์ที่เหมาะสมกับอุตสาหกรรมสร้างสรรค์ไทย โดยประสานความร่วมมือกับสถาบันข้อมูลขนาดใหญ่ (BDI) และหน่วยงานที่เกี่ยวข้อง',
    '๔.': '๔. ส่งเสริมการพัฒนากำลังคน Talent Pipeline ด้าน AI Creative ผ่านหลักสูตรการศึกษา การฝึกอบรม และ Sandbox สำหรับผู้ประกอบการ รวมถึงผลักดันการจัดทำมาตรฐานวิชาชีพ จรรยาบรรณ และกรอบกฎหมายด้านลิขสิทธิ์ผลงาน AI-generated',
    '๕.': '๕. ส่งเสริมการส่งออกผลงานและเครื่องมือ AI Creative ของไทยสู่ตลาดอาเซียนและระดับสากล ผ่านความร่วมมือกับกรมส่งเสริมการค้าระหว่างประเทศ (DITP) และงาน Bangkok International Content Market',
}

duty_count = 0
for i in range(sec_start, sec_end):
    para = doc.paragraphs[i]
    txt = para.text.strip()
    if txt in duty_replacements:
        set_para_text(para, duty_replacements[txt])
        duty_count += 1

print(f'  ✓ Filled {duty_count} duties (๒-๕)')

doc.save(OUTPUT)
print(f'\n✅ Saved: {OUTPUT}')

# ════════════════ VERIFY ════════════════
print('\n=== VERIFY Section 9 ===')
doc = Document(OUTPUT)
sec_start, sec_end = find_section_bounds(doc, 'อุตสาหกรรมสร้างสรรค์')
for i in range(sec_start, sec_end):
    txt = doc.paragraphs[i].text.strip()
    if txt:
        print(f'  [{i}] {txt[:100]}')
