# -*- coding: utf-8 -*-
"""สร้างเอกสาร Word (แก้ไขได้) — COE Creative Economy ตามรูปแบบ BDI Med-AI (ไฟล์ 57)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = 'TH Sarabun New'
doc = Document()

# default style
st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(15)
st._element.rPr.rFonts.set(qn('w:cs'), FONT)
st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def set_cs(run):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:cs'), FONT); rf.set(qn('w:ascii'), FONT); rf.set(qn('w:hAnsi'), FONT)

def para(text='', bold=False, size=15, align=None, color=None, space_after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    set_cs(r)
    return p

def head(text):  # section heading (bold)
    return para(text, bold=True, size=15, space_after=3)

def bullet(text, lvl=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + lvl*0.6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(15); set_cs(r)
    return p

BLUE = (0x1F, 0x4E, 0x78)

# ---- Title ----
para('Center of Excellence', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, space_after=2)
para('เพื่อเสนอภายใต้แผนยุทธศาสตร์ปัญญาประดิษฐ์ของประเทศไทย', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, space_after=10)

# (1)
head('(1) ชื่อโครงการ/แผนงาน')
para('ศูนย์ความเป็นเลิศด้านปัญญาประดิษฐ์เพื่อเศรษฐกิจสร้างสรรค์ (AI Center of Excellence for Creative Economy)', bold=True)

# (2)
head('(2) ความสำคัญ/หลักการและเหตุผล')
para('ประเทศไทยกำลังก้าวสู่การประยุกต์ใช้ AI อย่างเป็นระบบ โดยอุตสาหกรรมสร้างสรรค์ (ภาพยนตร์ แอนิเมชัน เกม เพลง คอนเทนต์ โฆษณา ออกแบบ) เป็นหนึ่งในเป้าหมายสำคัญที่ AI สามารถเป็นตัวคูณผลิตภาพ ลดต้นทุนการผลิต เพิ่มความเร็ว และเปิดตลาดส่งออกได้ทันที ปัจจุบันเศรษฐกิจสร้างสรรค์คิดเป็นร้อยละ 8.78 ของ GDP (รายงาน CEA พฤษภาคม 2569) ซึ่งหากใส่ AI เป็นตัวเร่ง คือการอัปเกรดทั้งฐาน มิใช่สร้างใหม่จากศูนย์')
para('อย่างไรก็ดี ระบบนิเวศ AI สำหรับงานสร้างสรรค์ไทยยังขาดกลไกที่เชื่อมทั้งวงจร ตั้งแต่ข้อมูล–การผลิต–การตรวจคุณภาพ–การรับรอง–ตลาด–การเทียบวัด แผนงานนี้จึงมุ่งจัดให้เกิดกลไกขับเคลื่อนระดับชาติที่ดำเนินงานในรูปเครือข่ายความร่วมมือ (mechanism & network) โดยใช้ประโยชน์จากงบประมาณและโปรแกรมที่หน่วยงานรัฐมีอยู่เดิม (leverage/co-fund) และสนับสนุนแบบจ่ายเมื่อเกิดผล (pay-for-outcome) มิใช่การจัดตั้งศูนย์ อาคาร หรือสถาบันขึ้นใหม่ ทั้งนี้เพื่อสร้างความน่าเชื่อถือ คุ้มครองสิทธิผู้สร้างสรรค์ในยุค AI และผลักดันให้คอนเทนต์/เครื่องมือ AI ไทยเข้าสู่ตลาดทั้งในและต่างประเทศ สอดคล้องกับแผนปฏิบัติการพัฒนาระบบนิเวศปัญญาประดิษฐ์แห่งชาติ (ทั้งสี่ยุทธศาสตร์ IGNITE/GROW/GOVERN/EQUIP) และวางไทยสู่การเป็นศูนย์กลางเศรษฐกิจสร้างสรรค์ที่ขับเคลื่อนด้วย AI ของอาเซียนอย่างยั่งยืน')

# (3)
head('(3) วัตถุประสงค์การจัดตั้ง Center of Excellence')
para('เพื่อให้ผู้ประกอบการและผู้สร้างสรรค์ไทยอย่างน้อย 5,000 ราย เข้าถึงและใช้ AI ในการผลิตได้จริง ลดต้นทุนการผลิตในสาขาเป้าหมายอย่างมีนัยสำคัญ และมีผลงาน/เครื่องมือ AI สร้างสรรค์ของไทยเข้าสู่ตลาดต่างประเทศได้ ภายในปี 2570')
para('โดยการขับเคลื่อนให้เกิดระบบนิเวศ AI เศรษฐกิจสร้างสรรค์ที่ครบวงจร (Creative AI Content Lifecycle) ครอบคลุม 6 มิติหลัก')
items3 = [
 ('1. การบริหารจัดการข้อมูลและสิทธิ์งานสร้างสรรค์ (Creative Data & Rights Management):',' จัดให้มีคลังสิทธิ์งานสร้างสรรค์และบัญชีความยินยอม (consent ledger) ที่ลิขสิทธิ์ถูกต้อง ป้อนชุดข้อมูลสร้างสรรค์ไทยเข้าสู่ ThaiLLM พร้อมกลไกจ่ายค่าตอบแทนคืนผู้สร้างสรรค์ (royalty) โดยใช้โครงสร้างพื้นฐานของ BDI/NECTEC ที่มีอยู่'),
 ('2. การผลิตด้วยปัญญาประดิษฐ์ (Creative AI Production):',' สนับสนุนผู้ประกอบการ/ครีเอเตอร์ให้ใช้เครื่องมือ AI ลดต้นทุนและผลิตงานจริง ผ่านกลไกคูปอง/เครดิต (leverage สสว. BDS, depa) และกลไกคืนเงินค่าผลิต (cash rebate ร่วม BOI) รวมถึงสายการผลิตพากย์/แปลคอนเทนต์ด้วย AI'),
 ('3. การประเมินคุณภาพ (Creative AI Validation — Thai-LQM):',' จัดให้มีตรารับรองคุณภาพและความถูกต้องเชิงภาษา-วัฒนธรรม สำหรับงานแปล/พากย์/ผลงานสร้างสรรค์ที่ใช้ AI โดยประสานราชบัณฑิตยสภาและ ETDA'),
 ('4. การรับรองที่มาและทรัพย์สินทางปัญญา (Provenance & IP — Trusted Thai AI):',' จัดให้มีระบบรับรองแหล่งที่มา (C2PA) ผ่านเกณฑ์ EU AI Act มาตรา 50 และทะเบียนทรัพย์สินทางปัญญา เพื่อให้คอนเทนต์ไทยผ่านด่านตลาดต่างประเทศ โดยประสาน ETDA สมอ. และกรมทรัพย์สินทางปัญญา'),
 ('5. การขยายผลสู่ตลาดและเชิงพาณิชย์ (Market & Commercialization):',' ผลักดันผลงาน/เครื่องมือ/IP สร้างสรรค์ของไทยเข้าสู่ตลาดและการส่งออก ผ่านกลไกของ DITP, BOI, EXIM และเวทีจับคู่ธุรกิจ รวมถึงเกม/อีสปอร์ตและการรับจ้างผลิต (service export)'),
 ('6. มาตรฐานและการเทียบวัด (Creative AI Benchmark):',' จัดทำชุดทดสอบและมาตรฐานคุณภาพงานสร้างสรรค์ AI ภาษา/วัฒนธรรมไทย เพื่อเป็นมาตรฐานอ้างอิงระดับชาติและภูมิภาค'),
]
for b,t in items3:
    p = doc.add_paragraph(style='List Number'); p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(3)
    r1=p.add_run(b); r1.bold=True; r1.font.size=Pt(15); set_cs(r1)
    r2=p.add_run(t); r2.font.size=Pt(15); set_cs(r2)

# (4)
head('(4) กลุ่มเป้าหมายและพื้นที่ที่ดำเนินการ')
for t in ['ผู้ประกอบการ/วิสาหกิจสร้างสรรค์: สตูดิโอภาพยนตร์/แอนิเมชัน/VFX สตูดิโอเกม ค่ายเพลง เอเจนซีโฆษณา นักออกแบบ และวิสาหกิจขนาดกลางและขนาดย่อม',
          'ผู้พัฒนาเทคโนโลยีสร้างสรรค์ (CreativeTech): สตาร์ทอัพและผู้พัฒนาเครื่องมือ AI สำหรับงานสร้างสรรค์',
          'ผู้สร้างสรรค์/บุคลากร และเศรษฐกิจฐานราก: ครีเอเตอร์ ฟรีแลนซ์ บุคลากรในอุตสาหกรรม และผู้ประกอบการชุมชน (OTOP/วิสาหกิจชุมชน)']:
    bullet(t)

# (5)
head('(5) หน่วยงานที่เกี่ยวข้อง/หน่วยงานที่ร่วมดำเนินงาน ในช่วง 2 ปีแรก')
for t in ['สมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย (AICAT) — ผู้ประสานหลัก',
          'CEA · depa · NIA · BDI · NECTEC / สวทช.',
          'DITP · สสว. · BOI · EXIM',
          'ETDA · สมอ. · กรมทรัพย์สินทางปัญญา · สำนักงานราชบัณฑิตยสภา',
          'สมาคมวิชาชีพ: TACGA · DAAT · สมาคมผู้กำกับภาพยนตร์ไทย · สมาคมคอนเทนต์ครีเอเตอร์ไทย · MarTech · Thai IoT · DUGA · สมาคมสมาพันธ์โอเพนซอร์สแห่งประเทศไทย',
          'แหล่งทุนที่ร่วม leverage: กองทุนพัฒนาสื่อปลอดภัยและสร้างสรรค์ · กองทุนพัฒนาดิจิทัลฯ (DE Fund) · กองทุน ววน./บพข. · กองทุน กทปส.']:
    bullet(t)

# (6)
head('(6) งบประมาณโครงการในช่วง 2 ปีแรก — งบประมาณรวม 360 ล้านบาท')
para('(ตั้งไว้ระดับนำร่อง — ปรับได้ เทียบ CoE สุขภาพ 500 ลบ./2 ปี · งบ COE ก้อนนี้ leverage งบหน่วยงานเดิมและเอกชนอีกหลายเท่า)', size=13)
budget = [
 ('1','Creative Data & Rights Management (คลังสิทธิ์+corpus→ThaiLLM+royalty)','40','40'),
 ('2','Creative AI Production (voucher/credit + cash rebate + dub-thai)','55','55'),
 ('3','Creative AI Validation — Thai-LQM (ตรารับรองคุณภาพ-วัฒนธรรม)','15','15'),
 ('4','Provenance & IP — Trusted Thai AI (C2PA + EU AI Act ม.50 + ทะเบียน IP)','35','35'),
 ('5','Market & Commercialization (export/licensing/matching/game-esports)','25','25'),
 ('6','Creative AI Benchmark (มาตรฐานคุณภาพงานสร้างสรรค์ AI)','10','10'),
]
t6 = doc.add_table(rows=1, cols=4); t6.style='Table Grid'
hdr=['','กิจกรรม (โปรแกรมหลัก)','ปีที่ 1 (ลบ.)','ปีที่ 2 (ลบ.)']
for i,h in enumerate(hdr):
    c=t6.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(14); set_cs(r)
for no,name,y1,y2 in budget:
    row=t6.add_row().cells
    for i,v in enumerate([no,name,y1,y2]):
        row[i].text=''; r=row[i].paragraphs[0].add_run(v); r.font.size=Pt(14); set_cs(r)
row=t6.add_row().cells
for i,v in enumerate(['','รวม','180','180']):
    row[i].text=''; r=row[i].paragraphs[0].add_run(v); r.bold=True; r.font.size=Pt(14); set_cs(r)

# (7)
head('(7) ผลผลิต (Output) ภายใน 2 ปี')
for t in ['ผู้ประกอบการ/ครีเอเตอร์ใช้เครื่องมือ AI ในการผลิตได้จริงอย่างน้อย 5,000 ราย และมีข้อมูลฐานต้นทุนการผลิตเทียบก่อน-หลัง',
          'คลังสิทธิ์งานสร้างสรรค์ + ชุดข้อมูล (corpus) ลิขสิทธิ์สะอาด เปิดให้ใช้งานและป้อนเข้า ThaiLLM',
          'ตรารับรองคุณภาพ (Thai-LQM) และตรารับรองที่มา (Trusted Thai AI) ที่ใช้งานจริง ผ่านเกณฑ์ตลาดต่างประเทศ (EU AI Act ม.50)',
          'ผลงาน/เครื่องมือ/IP สร้างสรรค์ของไทยที่เข้าสู่ตลาดและส่งออก พร้อมมูลค่าดีลที่วัดได้']:
    bullet(t)

# (8)
head('(8) ผลกระทบ/ประโยชน์ที่คาดว่าจะได้')
for t in ['ร่วมขับเคลื่อนสัดส่วนเศรษฐกิจสร้างสรรค์ใน GDP จากร้อยละ 8.78 สู่ร้อยละ 12 (วิสัยทัศน์ระดับชาติ — มิใช่ตัวชี้วัดผูกพันรายโครงการ)',
          'ต้นทุนการผลิตคอนเทนต์ในสาขาเป้าหมายลดลงประมาณร้อยละ 30 (วัดด้วย Cost Index)',
          'ผู้สร้างสรรค์ไทยได้รับความคุ้มครองลิขสิทธิ์ในยุค AI และมีรายได้จากการใช้ข้อมูล (royalty)',
          'ดึงเม็ดเงินการผลิต/การส่งออก และการจ้างงานทักษะสูงเข้าประเทศ ผ่านกลไกจ่ายเมื่อเกิดผล',
          'ยกระดับขีดความสามารถผู้ประกอบการไทยและอุตสาหกรรม CreativeTech สู่การเป็นศูนย์กลางอาเซียน']:
    bullet(t)

# (9)
head('(9) กลไกการดำเนินงานและการกำกับ')
para('ดำเนินงานในรูปคณะทำงาน + คณะอนุทำงาน + สำนักงานเลขานุการ/PMO (เครือข่ายเสมือน ไม่ใช่หน่วยงาน/อาคารใหม่) โดย AICAT เป็นผู้ประสาน ใช้กลไกสนับสนุนแบบ co-pay/matching/rebate/first-loss/pay-for-outcome บนงบและโปรแกรมที่หน่วยงานเจ้าภาพมีอยู่ ติดตามและรายงานผลต่อคณะอนุกรรมการขับเคลื่อนแผนด้าน AI แห่งชาติเป็นระยะ โดยวัดผลเป็นมูลค่าทางการเงิน จำนวนดีล/สัญญา จำนวนราย และจำนวนชิ้นงาน')

# (10)
head('(10) แผนกิจกรรมในช่วง 2 ปีแรก')
plan = ['1. Creative Data & Rights Management','2. Creative AI Production (voucher/rebate/dub-thai)','3. Creative AI Validation — Thai-LQM','4. Provenance & IP — Trusted Thai AI','5. Market & Commercialization','6. Creative AI Benchmark']
t10 = doc.add_table(rows=1, cols=3); t10.style='Table Grid'
for i,h in enumerate(['กิจกรรม (โปรแกรมหลัก)','ปีที่ 1','ปีที่ 2']):
    c=t10.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(14); set_cs(r)
for name in plan:
    row=t10.add_row().cells
    r=row[0].paragraphs[0].add_run(name); r.font.size=Pt(14); set_cs(r)
    for i in (1,2):
        row[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=row[i].paragraphs[0].add_run('✔'); r.font.size=Pt(14); set_cs(r)

para('')
para('จัดทำโดย สมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย (AICAT) ในฐานะผู้ประสานคณะทำงานฯ · มิถุนายน 2569 · ฉบับร่างเพื่อทบทวน', size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE)

doc.save('57_COE_Creative_Economy_BDI-format.docx')
print('saved 57_COE_Creative_Economy_BDI-format.docx')
