"""สร้าง Word + PDF เฉพาะ Section ๙ (standalone)
สำหรับส่ง สดช. แยกออกมา ให้เค้าจัด format เอง
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import os

OUTPUT_DOCX = '/Users/maxvulcanx/Desktop/AI Project/COE AI Creative Economy/22_Section9_Standalone.docx'
OUTPUT_PDF_DIR = '/Users/maxvulcanx/Desktop/AI Project/COE AI Creative Economy'

# Font
THAI_FONT = 'TH SarabunPSK'  # Standard Thai government font
FALLBACK_FONT = 'Sarabun'

def set_thai_font(run, size_pt=16, bold=False):
    """Set Thai-compatible font."""
    run.font.name = THAI_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    # Set complex script font (for Thai)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), THAI_FONT)
    rFonts.set(qn('w:hAnsi'), THAI_FONT)
    rFonts.set(qn('w:cs'), THAI_FONT)

def add_para(doc, text, size=16, bold=False, align=None, space_after=6):
    """Add a paragraph with Thai font."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_thai_font(run, size_pt=size, bold=bold)
    return p

# ════════════════ Create Document ════════════════
doc = Document()

# Set page margins (A4 + reasonable margins)
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.0)

# Title
title = add_para(doc, '๙. คณะทำงานขับเคลื่อนศูนย์ความเป็นเลิศปัญญาประดิษฐ์',
                 size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
subtitle = add_para(doc, 'ด้านอุตสาหกรรมสร้างสรรค์ (Creative Economy)',
                    size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# ── องค์ประกอบ ──
add_para(doc, 'องค์ประกอบ', size=16, bold=True, space_after=8)

members = [
    ('๙.๑', 'ผู้แทนสมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย', 'ประธานคณะทำงาน'),
    ('๙.๒', 'ผู้ทรงคุณวุฒิ (๑) ผู้แทนสมาคมผู้กำกับภาพยนตร์ไทย', 'คณะทำงาน'),
    ('๙.๓', 'ผู้ทรงคุณวุฒิ (๒) ผู้แทนสมาคมผู้ประกอบการแอนิเมชันและคอมพิวเตอร์กราฟิกไทย', 'คณะทำงาน'),
    ('๙.๔', 'ผู้ทรงคุณวุฒิ (๓) ผู้แทนสมาคมโฆษณาดิจิทัล (ประเทศไทย)', 'คณะทำงาน'),
    ('๙.๕', 'ผู้ทรงคุณวุฒิ (๔) ผู้แทนสมาคมคอนเทนต์ครีเอเตอร์ไทย', 'คณะทำงาน'),
    ('๙.๖', 'ผู้ทรงคุณวุฒิ (๕) ด้านกฎหมายทรัพย์สินทางปัญญาและปัญญาประดิษฐ์', 'คณะทำงาน'),
    ('๙.๗', 'ผู้แทนสำนักงานส่งเสริมเศรษฐกิจดิจิทัล', 'คณะทำงาน'),
    ('๙.๘', 'ผู้แทนกรมส่งเสริมวัฒนธรรม กระทรวงวัฒนธรรม', 'คณะทำงาน'),
    ('๙.๙', 'ผู้แทนสำนักงานส่งเสริมเศรษฐกิจสร้างสรรค์ (องค์การมหาชน)', 'คณะทำงาน'),
    ('๙.๑๐', 'ผู้แทนกรมทรัพย์สินทางปัญญา กระทรวงพาณิชย์', 'คณะทำงาน'),
    ('๙.๑๑', 'ผู้แทนสำนักงานนวัตกรรมแห่งชาติ (องค์การมหาชน)', 'คณะทำงาน'),
    ('๙.๑๒', 'ผู้แทนสถาบันข้อมูลขนาดใหญ่ (องค์การมหาชน)', 'คณะทำงาน'),
    ('๙.๑๓', 'ผู้แทนกรมส่งเสริมการค้าระหว่างประเทศ กระทรวงพาณิชย์', 'คณะทำงาน'),
    ('๙.๑๔', 'ผู้แทนสมาคมสมาพันธ์โอเพนซอร์สแห่งประเทศไทย', 'คณะทำงาน'),
    ('๙.๑๕', 'ผู้แทนสมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย', 'คณะทำงานและเลขานุการ'),
    ('๙.๑๖', 'ผู้แทนสำนักงานนวัตกรรมแห่งชาติ (องค์การมหาชน)', 'คณะทำงานและผู้ช่วยเลขานุการ'),
    ('๙.๑๗', 'ผู้แทนสำนักงานคณะกรรมการดิจิทัลเพื่อเศรษฐกิจและสังคมแห่งชาติ', 'คณะทำงานและผู้ช่วยเลขานุการร่วม'),
]

# Use table for organized layout
member_table = doc.add_table(rows=len(members), cols=3)
member_table.autofit = False
member_table.allow_autofit = False
# Set column widths
widths = [Cm(1.5), Cm(11.5), Cm(4.5)]
for i, w in enumerate(widths):
    for cell in member_table.columns[i].cells:
        cell.width = w

for i, (num, position, role) in enumerate(members):
    row = member_table.rows[i]
    cells = row.cells
    # Num
    cells[0].text = ''
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(num)
    set_thai_font(run, size_pt=15)
    # Position
    cells[1].text = ''
    p = cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(position)
    set_thai_font(run, size_pt=15)
    # Role
    cells[2].text = ''
    p = cells[2].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(role)
    set_thai_font(run, size_pt=15)

# Spacing after table
add_para(doc, '', size=14, space_after=12)

# ── หน้าที่และอำนาจ ──
add_para(doc, 'หน้าที่และอำนาจ', size=16, bold=True, space_after=8)

duties = [
    'ศึกษา วิเคราะห์ และจัดทำข้อเสนอเชิงนโยบาย แผนปฏิบัติการ แนวทางการดำเนินงาน และแผนงาน/โครงการด้านปัญญาประดิษฐ์ ที่สอดคล้องกับแผนปฏิบัติการด้านปัญญาประดิษฐ์แห่งชาติ เพื่อเสนอต่อคณะอนุกรรมการพิจารณา',
    'การกำหนดทิศทางและนโยบายเชิงยุทธศาสตร์: ศึกษา วิเคราะห์ และจัดทำข้อเสนอเชิงนโยบาย ยุทธศาสตร์ และแผนงานด้านปัญญาประดิษฐ์สำหรับอุตสาหกรรมสร้างสรรค์ของประเทศ เพื่อยกระดับเศรษฐกิจสร้างสรรค์ของไทยและขับเคลื่อนสู่เวทีระดับสากล',
    'การจัดตั้งและดำเนินการศูนย์ความเป็นเลิศ: จัดตั้งและดำเนินการศูนย์ความเป็นเลิศปัญญาประดิษฐ์ด้านอุตสาหกรรมสร้างสรรค์ ทั้งในสถานที่ตั้งและในระบบดิจิทัล พร้อมศึกษาและผลักดันการพัฒนาให้เป็นสำนักงานหรือหน่วยงานเฉพาะในระยะถัดไป เพื่อเป็นศูนย์กลางในการประสาน พัฒนา เผยแพร่องค์ความรู้ และให้บริการแก่ผู้ประกอบการ ผู้สร้างสรรค์ และเครือข่ายที่เกี่ยวข้อง',
    'การบริหารจัดการระบบนิเวศ: เชื่อมโยงและบูรณาการความร่วมมือระหว่างสถาบันการศึกษา ภาคอุตสาหกรรมสร้างสรรค์ ผู้ประกอบการ ผู้สร้างสรรค์เนื้อหา ภาครัฐ และเครือข่ายระดับสากล เพื่อสร้างระบบนิเวศที่เอื้อต่อการประยุกต์ใช้ปัญญาประดิษฐ์',
    'การพัฒนาโครงสร้างพื้นฐานและชุดข้อมูล: ส่งเสริมและสนับสนุนการสร้างชุดข้อมูล โครงสร้างพื้นฐานข้อมูล และเครื่องมือปัญญาประดิษฐ์ที่เหมาะสมกับอุตสาหกรรมสร้างสรรค์ไทย โดยประสานความร่วมมือกับหน่วยงานที่เกี่ยวข้อง',
    'การพัฒนาโครงการต้นแบบและขยายผลเชิงพาณิชย์: พัฒนา สนับสนุน และขับเคลื่อนโครงการต้นแบบและกรณีการใช้งานจริงในอุตสาหกรรมสร้างสรรค์ ให้เกิดผลเป็นรูปธรรมและสามารถขยายผลเชิงพาณิชย์ได้',
    'การกำกับดูแล มาตรฐานวิชาชีพ และจริยธรรม: ผลักดันการจัดทำกรอบนโยบายทรัพย์สินทางปัญญาสำหรับผลงานที่สร้างจากปัญญาประดิษฐ์ มาตรฐานวิชาชีพ และจรรยาบรรณด้านปัญญาประดิษฐ์ในอุตสาหกรรมสร้างสรรค์ ภายใต้หลักธรรมาภิบาล',
    'การพัฒนากำลังคนและศักยภาพบุคลากร: ส่งเสริมการพัฒนาทักษะของผู้ประกอบการ ผู้สร้างสรรค์เนื้อหา และบุคลากรในอุตสาหกรรมสร้างสรรค์ ให้สามารถประยุกต์ใช้ปัญญาประดิษฐ์ได้อย่างมีประสิทธิภาพและต่อเนื่อง',
    'การส่งเสริมตลาดและการนำเสนอผลงาน: ส่งเสริมการนำผลงานและเครื่องมือปัญญาประดิษฐ์สำหรับอุตสาหกรรมสร้างสรรค์ของไทยเข้าสู่ตลาดทั้งในและต่างประเทศ ผ่านความร่วมมือกับหน่วยงานที่เกี่ยวข้องและเครือข่ายระดับสากล',
    'การบูรณาการและบริหารจัดการทรัพยากร: บูรณาการและบริหารจัดการทรัพยากรด้านปัญญาประดิษฐ์สำหรับอุตสาหกรรมสร้างสรรค์ ทั้งในด้านโครงสร้างพื้นฐาน ข้อมูล บุคลากร และงบประมาณที่เกี่ยวข้อง ให้เกิดประสิทธิภาพสูงสุด',
    'การบริหารจัดการทั่วไป: แต่งตั้งคณะทำงานย่อย ที่ปรึกษา หรือเชิญผู้เชี่ยวชาญจากภาคส่วนต่าง ๆ เพื่อสนับสนุนการดำเนินงานเฉพาะด้านได้ตามความเหมาะสม',
    'รายงานผลการดำเนินงาน รวมทั้งเสนอข้อเสนอเชิงนโยบาย ข้อเสนอแนะ และแนวทางการดำเนินการต่อคณะอนุกรรมการขับเคลื่อนแผนด้านปัญญาประดิษฐ์แห่งชาติเป็นระยะ',
    'ปฏิบัติหน้าที่อื่นใดตามที่ได้รับมอบหมาย',
]

thai_nums = ['๑', '๒', '๓', '๔', '๕', '๖', '๗', '๘', '๙', '๑๐', '๑๑', '๑๒', '๑๓']

for i, duty_text in enumerate(duties):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.8)

    # Number in front
    run = p.add_run(f'{thai_nums[i]}. ')
    set_thai_font(run, size_pt=15, bold=True)

    # Duty content
    run = p.add_run(duty_text)
    set_thai_font(run, size_pt=15)

# Save Word
doc.save(OUTPUT_DOCX)
print(f'✅ Word saved: {OUTPUT_DOCX}')

# ════════════════ Convert to PDF via LibreOffice ════════════════
print('\nConverting to PDF...')
result = subprocess.run(
    ['/opt/homebrew/bin/soffice', '--headless', '--convert-to', 'pdf',
     '--outdir', OUTPUT_PDF_DIR, OUTPUT_DOCX],
    capture_output=True, text=True, timeout=60
)

pdf_path = os.path.join(OUTPUT_PDF_DIR, '22_Section9_Standalone.pdf')
if os.path.exists(pdf_path):
    print(f'✅ PDF saved: {pdf_path}')
else:
    print(f'❌ PDF conversion failed')
    print(f'stdout: {result.stdout}')
    print(f'stderr: {result.stderr}')
