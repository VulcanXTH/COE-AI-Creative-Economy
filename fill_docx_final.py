"""ลง Section ๙ (Final v5) เข้า .docx ของ สดช.

Output: 21_คำสั่งแต่งตั้ง_Final.docx
"""
from docx import Document
from copy import deepcopy
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import sys

INPUT = '/Users/maxvulcanx/Downloads/Tham/690518 (ร่าง) คำสั่งแต่งตั้งคณะทำงาน COE Updated.docx'
OUTPUT = '/Users/maxvulcanx/Desktop/AI Project/COE AI Creative Economy/21_คำสั่งแต่งตั้ง_Final.docx'

def set_para_text(para, new_text):
    """Replace paragraph text, preserving first run formatting."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)

def insert_para_after(ref_para, new_text):
    """Insert paragraph after ref_para with same XML style. Returns new Paragraph."""
    new_p = deepcopy(ref_para._p)
    for t in new_p.iter(qn('w:t')):
        t.text = ''
    first_t = new_p.find('.//' + qn('w:t'))
    if first_t is not None:
        first_t.text = new_text
    ref_para._p.addnext(new_p)
    return Paragraph(new_p, ref_para._parent)

doc = Document(INPUT)

# Find Section 9
section_start = None
for i, para in enumerate(doc.paragraphs):
    if 'อุตสาหกรรมสร้างสรรค์' in para.text and 'คณะทำงานขับเคลื่อน' in para.text:
        section_start = i
        break

if section_start is None:
    print("ERROR: Section 9 not found")
    sys.exit(1)

print(f'Section 9 starts at paragraph {section_start}')

# ════════════════ STEP 1: Replace org slot text ════════════════
org_replacements = [
    ('๙.๑ ผู้แทนหน่วยงานที่รับผิดชอบ (รัฐ/เอกชน)',
     '๙.๑ ผู้แทนสมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย'),
    ('๙.๒ - ๙.๖ ผู้ทรงคุณวุฒิที่ประธานแต่งตั้ง จำนวนไม่เกิน ๕ ท่าน',
     '๙.๒ ผู้ทรงคุณวุฒิ (๑) ผู้แทนสมาคมผู้กำกับภาพยนตร์ไทย'),
    ('๙.๗ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๑)',
     '๙.๗ ผู้แทนสำนักงานส่งเสริมเศรษฐกิจดิจิทัล'),
    ('๙.๘ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๒)',
     '๙.๘ ผู้แทนกรมส่งเสริมวัฒนธรรม กระทรวงวัฒนธรรม'),
    ('๙.๙ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๓)',
     '๙.๙ ผู้แทนสำนักงานส่งเสริมเศรษฐกิจสร้างสรรค์ (องค์การมหาชน)'),
    ('๙.๑๐ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๔)',
     '๙.๑๐ ผู้แทนกรมทรัพย์สินทางปัญญา กระทรวงพาณิชย์'),
    ('๙.๑๑ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๕)',
     '๙.๑๑ ผู้แทนสำนักงานนวัตกรรมแห่งชาติ (องค์การมหาชน)'),
    ('๙.๑๒ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๖)',
     '๙.๑๒ ผู้แทนสถาบันข้อมูลขนาดใหญ่ (องค์การมหาชน)'),
    ('๙.๑๓ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๗)',
     '๙.๑๓ ผู้แทนกรมส่งเสริมการค้าระหว่างประเทศ กระทรวงพาณิชย์'),
    ('๙.๑๔ ผู้แทนหน่วยงานที่เกี่ยวข้อง (รัฐ/เอกชน ๘)',
     '๙.๑๔ ผู้แทนสมาคมสมาพันธ์โอเพนซอร์สแห่งประเทศไทย'),
    ('๙.๑๕ ผู้แทนหน่วยงานที่รับผิดชอบ',
     '๙.๑๕ ผู้แทนสมาคมสร้างสรรค์ปัญญาประดิษฐ์ไทย'),
    ('๙.๑๖ ผู้แทนหน่วยงานที่รับผิดชอบ',
     '๙.๑๖ ผู้แทนสำนักงานนวัตกรรมแห่งชาติ (องค์การมหาชน)'),
]

# Apply within Section 9 area (up to ~50 paragraphs after start)
search_range = min(section_start + 50, len(doc.paragraphs))
for i in range(section_start, search_range):
    para = doc.paragraphs[i]
    for old, new in org_replacements:
        if old in para.text:
            set_para_text(para, para.text.replace(old, new))
            break

# Remove "/สถาบันข้อมูลขนาดใหญ่" from 9.17 continuation line
for i in range(section_start, search_range):
    para = doc.paragraphs[i]
    if '/สถาบันข้อมูลขนาดใหญ่' in para.text:
        new_text = para.text.replace(' /สถาบันข้อมูลขนาดใหญ่', '').replace('/สถาบันข้อมูลขนาดใหญ่', '')
        set_para_text(para, new_text)
        print('  ✓ Removed /สถาบันข้อมูลขนาดใหญ่ from 9.17')
        break

print('✓ Org slots replaced')

# ════════════════ STEP 2: Insert ๙.๓-๙.๖ after ๙.๒ ════════════════
ref_para = None
for i in range(section_start, search_range):
    if doc.paragraphs[i].text.strip().startswith('๙.๒ ผู้ทรงคุณวุฒิ'):
        ref_para = doc.paragraphs[i]
        break

if ref_para:
    inserts = [
        '๙.๓ ผู้ทรงคุณวุฒิ (๒) ผู้แทนสมาคมผู้ประกอบการแอนิเมชันและคอมพิวเตอร์กราฟิกไทย\t\t\tคณะทำงาน',
        '๙.๔ ผู้ทรงคุณวุฒิ (๓) ผู้แทนสมาคมโฆษณาดิจิทัล (ประเทศไทย)\t\t\t\tคณะทำงาน',
        '๙.๕ ผู้ทรงคุณวุฒิ (๔) ผู้แทนสมาคมคอนเทนต์ครีเอเตอร์ไทย\t\t\t\tคณะทำงาน',
        '๙.๖ ผู้ทรงคุณวุฒิ (๕) ด้านกฎหมายทรัพย์สินทางปัญญาและปัญญาประดิษฐ์\t\tคณะทำงาน',
    ]
    last_para = ref_para
    for txt in inserts:
        last_para = insert_para_after(last_para, txt)
    print(f'✓ Inserted {len(inserts)} ผู้ทรงคุณวุฒิ paragraphs')

# Save intermediate + reload to refresh paragraph indices
doc.save(OUTPUT)
doc = Document(OUTPUT)

# Re-find section start
section_start = None
for i, para in enumerate(doc.paragraphs):
    if 'อุตสาหกรรมสร้างสรรค์' in para.text and 'คณะทำงานขับเคลื่อน' in para.text:
        section_start = i
        break

# ════════════════ STEP 3: Renumber original ๖ → ๑๒ and ๗ → ๑๓ ════════════════
for i in range(section_start, len(doc.paragraphs)):
    para = doc.paragraphs[i]
    txt = para.text.strip()
    # Stop at next section
    if txt.startswith('คณะทำงานขับเคลื่อน') and i > section_start:
        break
    if txt.startswith('๖. รายงานผลการดำเนินงาน'):
        new_text = para.text.replace('๖.', '๑๒.', 1)
        set_para_text(para, new_text)
        print('✓ Renumbered ๖ → ๑๒ (รายงาน)')
    elif txt.startswith('๗. ปฏิบัติหน้าที่อื่นใด'):
        new_text = para.text.replace('๗.', '๑๓.', 1)
        set_para_text(para, new_text)
        print('✓ Renumbered ๗ → ๑๓ (ปฏิบัติอื่น)')

# ════════════════ STEP 4: Fill empty duties ๒-๕ ════════════════
duty_replacements = {
    '๒.': '๒. การกำหนดทิศทางและนโยบายเชิงยุทธศาสตร์: ศึกษา วิเคราะห์ และจัดทำข้อเสนอเชิงนโยบาย ยุทธศาสตร์ และแผนงานด้านปัญญาประดิษฐ์สำหรับอุตสาหกรรมสร้างสรรค์ของประเทศ เพื่อยกระดับเศรษฐกิจสร้างสรรค์ของไทยและขับเคลื่อนสู่เวทีระดับสากล',
    '๓.': '๓. การจัดตั้งและดำเนินการศูนย์ความเป็นเลิศ: จัดตั้งและดำเนินการศูนย์ความเป็นเลิศปัญญาประดิษฐ์ด้านอุตสาหกรรมสร้างสรรค์ ทั้งในสถานที่ตั้งและในระบบดิจิทัล พร้อมศึกษาและผลักดันการพัฒนาให้เป็นสำนักงานหรือหน่วยงานเฉพาะในระยะถัดไป เพื่อเป็นศูนย์กลางในการประสาน พัฒนา เผยแพร่องค์ความรู้ และให้บริการแก่ผู้ประกอบการ ผู้สร้างสรรค์ และเครือข่ายที่เกี่ยวข้อง',
    '๔.': '๔. การบริหารจัดการระบบนิเวศ: เชื่อมโยงและบูรณาการความร่วมมือระหว่างสถาบันการศึกษา ภาคอุตสาหกรรมสร้างสรรค์ ผู้ประกอบการ ผู้สร้างสรรค์เนื้อหา ภาครัฐ และเครือข่ายระดับสากล เพื่อสร้างระบบนิเวศที่เอื้อต่อการประยุกต์ใช้ปัญญาประดิษฐ์',
    '๕.': '๕. การพัฒนาโครงสร้างพื้นฐานและชุดข้อมูล: ส่งเสริมและสนับสนุนการสร้างชุดข้อมูล โครงสร้างพื้นฐานข้อมูล และเครื่องมือปัญญาประดิษฐ์ที่เหมาะสมกับอุตสาหกรรมสร้างสรรค์ไทย โดยประสานความร่วมมือกับหน่วยงานที่เกี่ยวข้อง',
}

filled_count = 0
for i in range(section_start, len(doc.paragraphs)):
    para = doc.paragraphs[i]
    txt = para.text.strip()
    if txt.startswith('คณะทำงานขับเคลื่อน') and i > section_start:
        break
    if txt in duty_replacements:
        set_para_text(para, duty_replacements[txt])
        filled_count += 1

print(f'✓ Filled {filled_count} empty duties (๒-๕)')

# ════════════════ STEP 5: Insert new duties ๖-๑๑ after duty ๕ ════════════════
target_duty_5 = None
for i in range(section_start, len(doc.paragraphs)):
    para = doc.paragraphs[i]
    if para.text.strip().startswith('๕. การพัฒนาโครงสร้างพื้นฐาน'):
        target_duty_5 = para
        break

if target_duty_5:
    new_duties = [
        '๖. การพัฒนาโครงการต้นแบบและขยายผลเชิงพาณิชย์: พัฒนา สนับสนุน และขับเคลื่อนโครงการต้นแบบและกรณีการใช้งานจริงในอุตสาหกรรมสร้างสรรค์ ให้เกิดผลเป็นรูปธรรมและสามารถขยายผลเชิงพาณิชย์ได้',
        '๗. การกำกับดูแล มาตรฐานวิชาชีพ และจริยธรรม: ผลักดันการจัดทำกรอบนโยบายทรัพย์สินทางปัญญาสำหรับผลงานที่สร้างจากปัญญาประดิษฐ์ มาตรฐานวิชาชีพ และจรรยาบรรณด้านปัญญาประดิษฐ์ในอุตสาหกรรมสร้างสรรค์ ภายใต้หลักธรรมาภิบาล',
        '๘. การพัฒนากำลังคนและศักยภาพบุคลากร: ส่งเสริมการพัฒนาทักษะของผู้ประกอบการ ผู้สร้างสรรค์เนื้อหา และบุคลากรในอุตสาหกรรมสร้างสรรค์ ให้สามารถประยุกต์ใช้ปัญญาประดิษฐ์ได้อย่างมีประสิทธิภาพและต่อเนื่อง',
        '๙. การส่งเสริมตลาดและการนำเสนอผลงาน: ส่งเสริมการนำผลงานและเครื่องมือปัญญาประดิษฐ์สำหรับอุตสาหกรรมสร้างสรรค์ของไทยเข้าสู่ตลาดทั้งในและต่างประเทศ ผ่านความร่วมมือกับหน่วยงานที่เกี่ยวข้องและเครือข่ายระดับสากล',
        '๑๐. การบูรณาการและบริหารจัดการทรัพยากร: บูรณาการและบริหารจัดการทรัพยากรด้านปัญญาประดิษฐ์สำหรับอุตสาหกรรมสร้างสรรค์ ทั้งในด้านโครงสร้างพื้นฐาน ข้อมูล บุคลากร และงบประมาณที่เกี่ยวข้อง ให้เกิดประสิทธิภาพสูงสุด',
        '๑๑. การบริหารจัดการทั่วไป: แต่งตั้งคณะทำงานย่อย ที่ปรึกษา หรือเชิญผู้เชี่ยวชาญจากภาคส่วนต่าง ๆ เพื่อสนับสนุนการดำเนินงานเฉพาะด้านได้ตามความเหมาะสม',
    ]
    last_para = target_duty_5
    for d in new_duties:
        last_para = insert_para_after(last_para, d)
    print(f'✓ Inserted {len(new_duties)} new duties (๖-๑๑)')

# ════════════════ Save Final ════════════════
doc.save(OUTPUT)
print(f'\n✅ Saved: {OUTPUT}')

# Verify
print('\n=== VERIFY Section ๙ ===')
doc = Document(OUTPUT)
section_start = None
for i, para in enumerate(doc.paragraphs):
    if 'อุตสาหกรรมสร้างสรรค์' in para.text and 'คณะทำงานขับเคลื่อน' in para.text:
        section_start = i
        break

for i in range(section_start, min(section_start + 60, len(doc.paragraphs))):
    txt = doc.paragraphs[i].text.strip()
    if txt and txt.startswith('คณะทำงานขับเคลื่อน') and i > section_start:
        print(f'  --- Next section at {i} ---')
        break
    if txt:
        preview = txt[:100] + ('...' if len(txt) > 100 else '')
        print(f'  [{i}] {preview}')
